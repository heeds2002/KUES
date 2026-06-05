from flask import Flask, render_template, request, redirect, url_for, session, send_file
import os
import pickle
import pytesseract
import PyPDF2
import sqlite3
from io import BytesIO
from docx import Document
from werkzeug.utils import secure_filename

from features import (
    extract_basic_features,
    generate_feature_explanation,
    adjust_ai_score,
    highlight_suspicious_text,
    extract_answer_text
)

from ocr_utils import extract_handwritten_text


app = Flask(__name__)
app.secret_key = "secretkey"

UPLOAD_FOLDER = "uploads"
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

model = pickle.load(open("model/ai_model.pkl", "rb"))
vectorizer = pickle.load(open("model/vectorizer.pkl", "rb"))

latest_report_data = {}


# =========================
# DATABASE INITIALIZATION
# =========================

def init_db():
    conn = sqlite3.connect("database.db")
    cursor = conn.cursor()

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE,
            password TEXT,
            role TEXT DEFAULT 'examiner'
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS results (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            filename TEXT,
            assessment_type TEXT,
            tested_record TEXT,
            reg_number TEXT,
            cat1_marks REAL DEFAULT 0,
            cat2_marks REAL DEFAULT 0,
            total_marks REAL DEFAULT 0,
            date TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)

    cursor.execute("SELECT * FROM users WHERE username = ?", ("admin",))
    admin = cursor.fetchone()

    if admin is None:
        cursor.execute("""
            INSERT INTO users (username, password, role)
            VALUES (?, ?, ?)
        """, ("admin", "admin123", "admin"))

    conn.commit()
    conn.close()


init_db()


# =========================
# HELPER FUNCTIONS
# =========================

def save_record(filename, assessment_type, tested_record):
    conn = sqlite3.connect("database.db")
    cursor = conn.cursor()

    cursor.execute("""
        INSERT INTO results (
            filename,
            assessment_type,
            tested_record
        )
        VALUES (?, ?, ?)
    """, (filename, assessment_type, tested_record))

    conn.commit()
    conn.close()


def extract_text_from_files(files, label):
    extracted_text = ""
    uploaded_filenames = []

    for file in files:
        if file.filename == "":
            continue

        filename = secure_filename(file.filename)
        uploaded_filenames.append(filename)

        filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
        file.save(filepath)

        try:
            extracted_text += f"\n--- {label}: {filename} ---\n"

            if filename.lower().endswith((".png", ".jpg", ".jpeg")):
                extracted_text += extract_handwritten_text(filepath) + "\n"

            elif filename.lower().endswith(".txt"):
                with open(filepath, "r", encoding="utf-8") as f:
                    extracted_text += f.read() + "\n"

            elif filename.lower().endswith(".pdf"):
                pdf = PyPDF2.PdfReader(filepath)

                for page_number, page in enumerate(pdf.pages, start=1):
                    page_text = page.extract_text()

                    if page_text:
                        extracted_text += f"\n--- Page {page_number} ---\n"
                        extracted_text += page_text + "\n"

            elif filename.lower().endswith(".docx"):
                doc = Document(filepath)

                for para in doc.paragraphs:
                    extracted_text += para.text + "\n"

            else:
                extracted_text += f"\nUnsupported file format: {filename}\n"

        except Exception as e:
            extracted_text += f"\nError extracting from {filename}: {e}\n"

    return extracted_text, uploaded_filenames


def analyze_text(text):
    answer_text = extract_answer_text(text)

    features = extract_basic_features(answer_text)
    reasons = generate_feature_explanation(features)
    highlighted_text = highlight_suspicious_text(answer_text, features)

    ai_score = 0
    human_score = 0
    original_ml_score = 0
    ai_result = "Unknown"

    breakdown = {
        "increased": [],
        "reduced": []
    }

    if answer_text.strip() != "":
        text_vector = vectorizer.transform([answer_text])
        prediction = model.predict_proba(text_vector)[0]

        ml_ai_score = prediction[1] * 100
        original_ml_score = ml_ai_score

        ai_score, breakdown = adjust_ai_score(ml_ai_score, features)
        human_score = 100 - ai_score

        ai_result = "AI Written" if ai_score > human_score else "Human Written"

    return {
        "answer_text": answer_text,
        "features": features,
        "reasons": reasons,
        "highlighted_text": highlighted_text,
        "ai_score": ai_score,
        "human_score": human_score,
        "original_ml_score": original_ml_score,
        "ai_result": ai_result,
        "breakdown": breakdown
    }


# =========================
# AUTH ROUTES
# =========================

@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]

        conn = sqlite3.connect("database.db")
        cursor = conn.cursor()

        try:
            cursor.execute("""
                INSERT INTO users (username, password, role)
                VALUES (?, ?, ?)
            """, (username, password, "examiner"))

            conn.commit()
            conn.close()

            return redirect(url_for("login"))

        except sqlite3.IntegrityError:
            conn.close()
            return "Username already exists. Choose another username."

    return render_template("register.html")


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]

        conn = sqlite3.connect("database.db")
        cursor = conn.cursor()

        cursor.execute("""
            SELECT * FROM users
            WHERE username = ?
            AND password = ?
        """, (username, password))

        user = cursor.fetchone()
        conn.close()

        if user:
            session["user"] = username
            session["role"] = user[3]

            if user[3] == "admin":
                return redirect(url_for("admin"))

            return redirect(url_for("home"))

        return "Invalid Credentials"

    return render_template("login.html")


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


# =========================
# MAIN ROUTES
# =========================

@app.route("/")
def home():
    if "user" not in session:
        return redirect(url_for("login"))

    return render_template("index.html")


@app.route("/upload", methods=["GET", "POST"])
def upload():
    global latest_report_data

    if "user" not in session:
        return redirect(url_for("login"))

    if request.method == "GET":
        return render_template("upload.html")

    cat1_files = request.files.getlist("cat1_files")
    cat2_files = request.files.getlist("cat2_files")

    if not cat1_files or all(file.filename == "" for file in cat1_files):
        return "Please upload CAT 1 file(s)."

    cat1_text, cat1_names = extract_text_from_files(cat1_files, "CAT 1")
    cat1_analysis = analyze_text(cat1_text)

    all_extracted_text = "\n\n===== CAT 1 EXTRACTED TEXT =====\n" + cat1_text
    all_highlighted_text = "\n\n===== CAT 1 ANSWER TEXT =====\n" + cat1_analysis["highlighted_text"]

    tested_record_parts = [
        f"CAT 1: {cat1_analysis['ai_result']} ({round(cat1_analysis['ai_score'], 2)}%)"
    ]

    all_filenames = cat1_names

    final_ai_score = cat1_analysis["ai_score"]
    final_human_score = cat1_analysis["human_score"]
    final_original_ml_score = cat1_analysis["original_ml_score"]
    final_ai_result = cat1_analysis["ai_result"]
    final_features = cat1_analysis["features"]
    final_reasons = cat1_analysis["reasons"]
    final_breakdown = cat1_analysis["breakdown"]

    cat2_analysis = None

    if cat2_files and not all(file.filename == "" for file in cat2_files):
        cat2_text, cat2_names = extract_text_from_files(cat2_files, "CAT 2")
        cat2_analysis = analyze_text(cat2_text)

        all_extracted_text += "\n\n===== CAT 2 EXTRACTED TEXT =====\n" + cat2_text
        all_highlighted_text += "\n\n===== CAT 2 ANSWER TEXT =====\n" + cat2_analysis["highlighted_text"]

        tested_record_parts.append(
            f"CAT 2: {cat2_analysis['ai_result']} ({round(cat2_analysis['ai_score'], 2)}%)"
        )

        all_filenames += cat2_names

        final_ai_score = (
            cat1_analysis["ai_score"] + cat2_analysis["ai_score"]
        ) / 2

        final_human_score = 100 - final_ai_score

        final_original_ml_score = (
            cat1_analysis["original_ml_score"] + cat2_analysis["original_ml_score"]
        ) / 2

        final_ai_result = "AI Written" if final_ai_score > final_human_score else "Human Written"

        final_reasons = cat1_analysis["reasons"] + cat2_analysis["reasons"]

        final_breakdown = {
            "increased": (
                cat1_analysis["breakdown"]["increased"]
                + cat2_analysis["breakdown"]["increased"]
            ),
            "reduced": (
                cat1_analysis["breakdown"]["reduced"]
                + cat2_analysis["breakdown"]["reduced"]
            )
        }

    combined_filename = ", ".join(all_filenames)
    tested_record = " | ".join(tested_record_parts)

    save_record(
        combined_filename,
        "CAT 1 and CAT 2" if cat2_analysis else "CAT 1",
        tested_record
    )

    latest_report_data = {
        "ai_result": final_ai_result,
        "ai_score": final_ai_score,
        "human_score": final_human_score,
        "original_ml_score": final_original_ml_score,
        "features": final_features,
        "reasons": final_reasons,
        "breakdown": final_breakdown,
        "highlighted_text": all_highlighted_text,
        "extracted_text": all_extracted_text,
        "tested_record": tested_record
    }

    return render_template(
        "result.html",
        ai_score=final_ai_score,
        human_score=final_human_score,
        original_ml_score=final_original_ml_score,
        ai_result=final_ai_result,

        cat1_score=cat1_analysis["ai_score"],
        cat1_human_score=cat1_analysis["human_score"],
        cat1_result=cat1_analysis["ai_result"],

        cat2_score=cat2_analysis["ai_score"] if cat2_analysis else 0,
        cat2_human_score=cat2_analysis["human_score"] if cat2_analysis else 0,
        cat2_result=cat2_analysis["ai_result"] if cat2_analysis else "Not Uploaded",
        has_cat2=True if cat2_analysis else False,

        extracted_text=all_extracted_text,
        highlighted_text=all_highlighted_text,
        features=final_features,
        reasons=final_reasons,
        breakdown=final_breakdown
    )


# =========================
# RECORDS ROUTES
# =========================

@app.route("/records")
def records():
    if "user" not in session:
        return redirect(url_for("login"))

    conn = sqlite3.connect("database.db")
    cursor = conn.cursor()

    cursor.execute("SELECT * FROM results ORDER BY date DESC")
    records = cursor.fetchall()

    conn.close()

    return render_template("records.html", records=records)


@app.route("/update_marks/<int:record_id>", methods=["POST"])
def update_marks(record_id):
    if "user" not in session:
        return redirect(url_for("login"))

    reg_number = request.form.get("reg_number", "")
    cat1_marks = float(request.form.get("cat1_marks", 0) or 0)
    cat2_marks = float(request.form.get("cat2_marks", 0) or 0)

    total_marks = cat1_marks + cat2_marks

    conn = sqlite3.connect("database.db")
    cursor = conn.cursor()

    cursor.execute("""
        UPDATE results
        SET reg_number = ?,
            cat1_marks = ?,
            cat2_marks = ?,
            total_marks = ?
        WHERE id = ?
    """, (reg_number, cat1_marks, cat2_marks, total_marks, record_id))

    conn.commit()
    conn.close()

    return redirect(url_for("records"))


# =========================
# ADMIN ROUTES
# =========================

@app.route("/admin")
def admin():
    if "user" not in session or session.get("role") != "admin":
        return redirect(url_for("login"))

    conn = sqlite3.connect("database.db")
    cursor = conn.cursor()

    cursor.execute("SELECT * FROM users")
    users = cursor.fetchall()

    cursor.execute("SELECT * FROM results ORDER BY date DESC")
    records = cursor.fetchall()

    conn.close()

    return render_template("admin.html", users=users, records=records)


@app.route("/delete_user/<int:user_id>")
def delete_user(user_id):
    if "user" not in session or session.get("role") != "admin":
        return redirect(url_for("login"))

    conn = sqlite3.connect("database.db")
    cursor = conn.cursor()

    cursor.execute("SELECT username FROM users WHERE id = ?", (user_id,))
    user = cursor.fetchone()

    if user and user[0] != "admin":
        cursor.execute("DELETE FROM users WHERE id = ?", (user_id,))
        conn.commit()

    conn.close()

    return redirect(url_for("admin"))


@app.route("/edit_user/<int:user_id>", methods=["GET", "POST"])
def edit_user(user_id):
    if "user" not in session or session.get("role") != "admin":
        return redirect(url_for("login"))

    conn = sqlite3.connect("database.db")
    cursor = conn.cursor()

    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]
        role = request.form["role"]

        cursor.execute("""
            UPDATE users
            SET username = ?, password = ?, role = ?
            WHERE id = ?
        """, (username, password, role, user_id))

        conn.commit()
        conn.close()

        return redirect(url_for("admin"))

    cursor.execute("SELECT * FROM users WHERE id = ?", (user_id,))
    user = cursor.fetchone()

    conn.close()

    return render_template("edit_user.html", user=user)


@app.route("/delete_record/<int:record_id>")
def delete_record(record_id):
    if "user" not in session or session.get("role") != "admin":
        return redirect(url_for("login"))

    conn = sqlite3.connect("database.db")
    cursor = conn.cursor()

    cursor.execute("DELETE FROM results WHERE id = ?", (record_id,))

    conn.commit()
    conn.close()

    return redirect(url_for("admin"))


# =========================
# WORD REPORT ROUTE
# =========================

@app.route("/download_report")
def download_report():
    if "user" not in session:
        return redirect(url_for("login"))

    if not latest_report_data:
        return "No report available. Please analyze a document first."

    doc = Document()

    doc.add_heading("KABIANGA UNIVERSITY EXAMINER SYSTEM", 0)
    doc.add_heading("AI Detection Analysis Report", level=1)

    doc.add_heading("Final Results", level=2)

    doc.add_paragraph(f"Final Result: {latest_report_data['ai_result']}")
    doc.add_paragraph(f"Original ML Score: {round(latest_report_data['original_ml_score'], 2)}%")
    doc.add_paragraph(f"Final Hybrid AI Score: {round(latest_report_data['ai_score'], 2)}%")
    doc.add_paragraph(f"Human Score: {round(latest_report_data['human_score'], 2)}%")
    doc.add_paragraph(f"Tested Record: {latest_report_data['tested_record']}")

    doc.add_heading("Writing Pattern Analysis", level=2)

    features = latest_report_data["features"]

    doc.add_paragraph(f"Sentence Count: {features['sentence_count']}")
    doc.add_paragraph(f"Word Count: {features['word_count']}")
    doc.add_paragraph(f"Average Sentence Length: {round(features['avg_sentence_length'], 2)}")
    doc.add_paragraph(f"Sentence Variation / Burstiness: {round(features['sentence_variance'], 2)}")
    doc.add_paragraph(f"Vocabulary Diversity: {round(features['vocab_diversity'], 2)}")
    doc.add_paragraph(f"Spelling Errors: {features['spelling_errors']}")
    doc.add_paragraph(f"Error Rate: {round(features['error_rate'] * 100, 2)}%")
    doc.add_paragraph(f"AI Phrase Count: {features['ai_phrase_count']}")
    doc.add_paragraph(f"Quote Count: {features['quote_count']}")
    doc.add_paragraph(f"Dash / Em-Dash Count: {features['dash_count']}")
    doc.add_paragraph(f"Parenthesis Count: {features['parenthesis_count']}")
    doc.add_paragraph(f"Colon Count: {features['colon_count']}")
    doc.add_paragraph(f"Comma Count: {features['comma_count']}")
    doc.add_paragraph(f"Full Stop Count: {features['fullstop_count']}")

    doc.add_heading("Reason for Detection", level=2)

    for reason in latest_report_data["reasons"]:
        doc.add_paragraph(reason, style="List Bullet")

    doc.add_heading("Suspicion Level Breakdown", level=2)

    doc.add_heading("Features Increasing AI Probability", level=3)
    for item in latest_report_data["breakdown"]["increased"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Features Reducing AI Probability", level=3)
    for item in latest_report_data["breakdown"]["reduced"]:
        doc.add_paragraph(item, style="List Bullet")

    clean_text = (
        latest_report_data["highlighted_text"]
        .replace('<span class="highlight-ai" title="AI transition phrase">', '')
        .replace('<span class="highlight-ai">', '')
        .replace('<span class="ai-label">AI Phrase</span>', '')
        .replace('</span>', '')
    )

    doc.add_heading("Answer Text Used for Testing", level=2)
    doc.add_paragraph(clean_text)

    doc.add_heading("Full Extracted Text", level=2)
    doc.add_paragraph(latest_report_data["extracted_text"])

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name="KUES_Analysis_Report.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


if __name__ == "__main__":
    app.run(debug=True)